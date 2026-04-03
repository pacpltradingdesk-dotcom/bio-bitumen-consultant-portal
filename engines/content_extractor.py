"""
Bio Bitumen Consulting — Document Content Extractor
=====================================================
Reads actual DOCX file content and extracts structured data for the dashboard.
Maps every document type to its dashboard module.
"""
import os
import json
from pathlib import Path

DOC_ROOT = Path(r"C:\Users\HP\Desktop\Bio Bitumen Full Working all document")
CACHE_FILE = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                           "data", "extracted_content.json")


def _read_docx_text(filepath):
    """Extract plain text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(filepath)
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text.append(" | ".join(row_text))
        return "\n".join(text)
    except Exception as e:
        return f"Error reading: {e}"


def _read_xlsx_sheets(filepath):
    """Extract sheet names and first rows from Excel."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True)
        data = {}
        for name in wb.sheetnames:
            ws = wb[name]
            rows = []
            for row in ws.iter_rows(min_row=1, max_row=min(20, ws.max_row), values_only=True):
                vals = [str(v) for v in row if v is not None]
                if vals:
                    rows.append(vals)
            data[name] = rows
        wb.close()
        return data
    except Exception as e:
        return {"error": str(e)}


# ── DOCUMENT TYPE → DASHBOARD MODULE MAPPING ─────────────────────────
# This maps each sub-folder in the plant structure to the dashboard module
# that should display its content.

DOC_MODULE_MAP = {
    "01_DPR": "dashboard",           # Dashboard Home + DPR Generator
    "02_Drawings": "engineering",     # Engineering Drawings tab
    "03_Engineering": "plant_design", # Plant Design tab
    "04_BOQ": "procurement",          # Procurement tab (equipment + costs)
    "05_Safety": "compliance",        # Compliance tab (HAZOP, Fire, EIA)
    "05_FOR_CONTRACTOR": "procurement",
    "06_Approvals": "compliance",     # Compliance tab (licenses, approvals)
    "06_FOR_MACHINERY_SUPPLIER": "procurement",
    "07_Execution": "plant_design",   # Plant Design (Gantt, HR, Inventory)
    "08_Financials": "financial",     # Financial Model tab
    "09_Commercial": "buyers",        # Buyers tab (export, market, pitch)
    "09_Technical_Documents": "plant_design",
    "10_Bank_KYC_Documents": "compliance",
    "11_Legal_Documents": "compliance",
    "12_Regulatory_Documents": "compliance",
    "13_Investor_Documents": "dpr",   # DPR Generator
    "14_Govt_Scheme_Docs": "compliance",
    "15_HR_Operations": "plant_design",
    "16_Commercial_Agreements": "buyers",
    "PDF_Final": "files",            # File System
}


def build_content_index():
    """Build an index mapping each document to its dashboard module with key content."""
    if os.path.exists(CACHE_FILE):
        try:
            with open(CACHE_FILE, "r") as f:
                cached = json.load(f)
            if len(cached) > 100:  # Only use cache if substantial
                return cached
        except Exception:
            pass

    index = {}

    # Process all 7 plant folders
    for plant_folder in sorted(DOC_ROOT.iterdir()):
        if not plant_folder.name.startswith("PLANT_"):
            continue

        cap_key = plant_folder.name  # e.g., PLANT_20MT_Day_INR_8Cr

        for section_folder in sorted(plant_folder.iterdir()):
            if not section_folder.is_dir():
                continue

            section = section_folder.name
            module = DOC_MODULE_MAP.get(section, "files")

            for filepath in sorted(section_folder.iterdir()):
                if not filepath.is_file():
                    continue

                ext = filepath.suffix.lower()
                rel_path = str(filepath.relative_to(DOC_ROOT))

                entry = {
                    "filename": filepath.name,
                    "path": rel_path,
                    "capacity": cap_key,
                    "section": section,
                    "module": module,
                    "extension": ext,
                    "size": filepath.stat().st_size,
                }

                # For key document types, extract a summary
                if ext == ".docx" and filepath.stat().st_size < 500000:  # < 500KB
                    try:
                        text = _read_docx_text(str(filepath))
                        # Store first 500 chars as summary
                        entry["summary"] = text[:500] if text else ""
                        entry["full_text_length"] = len(text)
                    except Exception:
                        entry["summary"] = ""

                index[rel_path] = entry

    # Process root-level files
    for filepath in DOC_ROOT.iterdir():
        if filepath.is_file() and filepath.suffix.lower() in (".docx", ".xlsx", ".pdf", ".pptx"):
            rel_path = filepath.name
            index[rel_path] = {
                "filename": filepath.name,
                "path": rel_path,
                "capacity": "general",
                "section": "root",
                "module": "files",
                "extension": filepath.suffix.lower(),
                "size": filepath.stat().st_size,
            }

    # Save cache
    try:
        os.makedirs(os.path.dirname(CACHE_FILE), exist_ok=True)
        with open(CACHE_FILE, "w") as f:
            json.dump(index, f, indent=2)
    except Exception:
        pass

    return index


def get_documents_for_module(module_name, capacity=None):
    """Get all documents mapped to a specific dashboard module."""
    index = build_content_index()
    results = []
    for path, entry in index.items():
        if entry["module"] == module_name:
            if capacity and entry.get("capacity", "") != "general":
                if capacity not in entry.get("capacity", ""):
                    continue
            results.append(entry)
    return results


def get_document_summary(filepath):
    """Get the text summary of a specific document."""
    index = build_content_index()
    entry = index.get(filepath, {})
    return entry.get("summary", "")


def count_by_module():
    """Count documents per dashboard module."""
    index = build_content_index()
    counts = {}
    for entry in index.values():
        module = entry.get("module", "other")
        counts[module] = counts.get(module, 0) + 1
    return counts


def count_by_capacity():
    """Count documents per plant capacity."""
    index = build_content_index()
    counts = {}
    for entry in index.values():
        cap = entry.get("capacity", "general")
        counts[cap] = counts.get(cap, 0) + 1
    return counts
