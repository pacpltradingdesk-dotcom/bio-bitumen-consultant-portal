"""
Bio Bitumen Consultant Portal — DOCX Customizer
Find-and-replace placeholders in Word documents with customer-specific details.
"""
import os
import copy
from docx import Document


# Common placeholder patterns found in the document suite
DEFAULT_PLACEHOLDERS = [
    "Customer Name", "[Customer Name]", "<<Customer Name>>",
    "CUSTOMER NAME", "INVESTOR NAME", "[Investor Name]",
    "[Client Name]", "CLIENT NAME", "Client Name",
    "[Company Name]", "COMPANY NAME", "Company Name",
    "[Party Name]", "PARTY NAME", "Party Name",
]


def get_default_replacements(customer, plant=None, cfg=None):
    """Build a COMPLETE replacement dict from customer + plant + cfg data.
    Auto-fills ALL placeholders in any form/application/letter/DPR."""
    replacements = {}

    # ── Customer / Client Info ──
    cust_name = customer.get("name", "") if customer else ""
    cust_company = customer.get("company", "") or cust_name

    for pattern in DEFAULT_PLACEHOLDERS:
        if "investor" in pattern.lower() or "client" in pattern.lower() or "customer" in pattern.lower():
            replacements[pattern] = cust_name
        elif "company" in pattern.lower() or "party" in pattern.lower():
            replacements[pattern] = cust_company

    replacements["[Customer Email]"] = customer.get("email", "") if customer else ""
    replacements["[Customer Phone]"] = customer.get("phone", "") if customer else ""
    replacements["[Customer State]"] = customer.get("state", "") if customer else ""
    replacements["[Customer City]"] = customer.get("city", "") if customer else ""
    replacements["[Customer Address]"] = customer.get("address", "") if customer else ""

    # ── Date ──
    from datetime import datetime, timezone, timedelta
    IST = timezone(timedelta(hours=5, minutes=30))
    today = datetime.now(IST).strftime("%d %B %Y")
    replacements["[DATE]"] = today
    replacements["[Date]"] = today
    replacements["<<DATE>>"] = today
    replacements["[YEAR]"] = datetime.now(IST).strftime("%Y")

    # ── Plant / Capacity Info ──
    if plant:
        replacements["[Capacity]"] = plant.get("label", "")
        replacements["[Investment]"] = f"Rs {plant.get('inv_cr', '')} Crore"
        replacements["[Location]"] = plant.get("location", "")

    # ── CFG — Full Project Data (auto-fills everything) ──
    if cfg:
        tpd = cfg.get("capacity_tpd", 20)
        state = cfg.get("state", "")
        inv = cfg.get("investment_cr", 8)

        # Project Identity
        replacements["[Project Name]"] = cfg.get("project_name", f"Bio-Bitumen Plant {tpd:.0f} TPD")
        replacements["[DPR Version]"] = cfg.get("dpr_version", "v1.0")
        replacements["[Prepared By]"] = cfg.get("prepared_by", "PPS Anantams Corporation")

        # Capacity & Plant
        replacements["[Capacity TPD]"] = f"{tpd:.0f}"
        replacements["[Capacity MT/Day]"] = f"{tpd:.0f} MT/Day"
        replacements["[Plant Capacity]"] = f"{tpd:.0f} MT/Day"
        replacements["[Working Days]"] = f"{cfg.get('working_days', 300)}"
        replacements["[Annual Production]"] = f"{tpd * cfg.get('working_days', 300) * 0.4:.0f} MT"

        # Location & Site
        replacements["[State]"] = state
        replacements["[City]"] = cfg.get("location", "")
        replacements["[Site Address]"] = cfg.get("site_address", "")
        replacements["[Plot Length]"] = f"{cfg.get('plot_length_m', 120)} m"
        replacements["[Plot Width]"] = f"{cfg.get('plot_width_m', 80)} m"
        replacements["[Plot Area]"] = f"{cfg.get('plot_length_m', 120) * cfg.get('plot_width_m', 80)} sqm"
        replacements["[Site Area Acres]"] = f"{cfg.get('site_area_acres', 2)} acres"
        replacements["[Pincode]"] = cfg.get("site_pincode", "")

        # Financial
        replacements["[Total Investment]"] = f"Rs {inv:.2f} Crore"
        replacements["[Investment Crore]"] = f"{inv:.2f}"
        replacements["[Investment Lac]"] = f"{inv * 100:.0f}"
        replacements["[Loan Amount]"] = f"Rs {cfg.get('loan_cr', inv * 0.6):.2f} Crore"
        replacements["[Equity Amount]"] = f"Rs {cfg.get('equity_cr', inv * 0.4):.2f} Crore"
        replacements["[Debt Percent]"] = f"{(1 - cfg.get('equity_ratio', 0.4)) * 100:.0f}%"
        replacements["[Equity Percent]"] = f"{cfg.get('equity_ratio', 0.4) * 100:.0f}%"
        replacements["[Interest Rate]"] = f"{cfg.get('interest_rate', 0.115) * 100:.1f}%"
        replacements["[EMI]"] = f"Rs {cfg.get('emi_lac_mth', 0):.2f} Lac/month"
        replacements["[ROI]"] = f"{cfg.get('roi_pct', 0):.1f}%"
        replacements["[IRR]"] = f"{cfg.get('irr_pct', 0):.1f}%"
        replacements["[DSCR]"] = f"{cfg.get('dscr_yr3', 0):.2f}x"
        replacements["[Break Even]"] = f"{cfg.get('break_even_months', 0)} months"
        replacements["[Payback Period]"] = f"{cfg.get('break_even_months', 0)} months"
        replacements["[Revenue Year 5]"] = f"Rs {cfg.get('revenue_yr5_lac', 0):.0f} Lac"
        replacements["[Monthly Profit]"] = f"Rs {cfg.get('monthly_profit_lac', 0):.1f} Lac"
        replacements["[Selling Price]"] = f"Rs {cfg.get('selling_price_per_mt', 35000):,}/MT"

        # Process / Technology
        replacements["[Technology]"] = "Biomass Pyrolysis (CSIR-CRRI Licensed)"
        replacements["[Bio Oil Yield]"] = f"{cfg.get('bio_oil_yield_pct', 32)}%"
        replacements["[Bio Char Yield]"] = f"{cfg.get('bio_char_yield_pct', 28)}%"
        replacements["[Syngas Yield]"] = f"{cfg.get('syngas_yield_pct', 22)}%"
        replacements["[Blend Ratio]"] = f"{cfg.get('bio_blend_pct', 20)}%"
        replacements["[Biomass Source]"] = cfg.get("biomass_source", "Rice Straw / Wheat Straw")

        # Utilities
        replacements["[Power kW]"] = f"{cfg.get('power_kw', 100)} kW"
        replacements["[Water KLD]"] = f"{max(5, int(tpd * 1.5))} KLD"
        replacements["[Staff Count]"] = f"{cfg.get('staff', 20)}"

        # Equipment (summary)
        try:
            from state_manager import calculate_boq
            boq = calculate_boq(tpd, process_id=cfg.get("process_id", 1))
            replacements["[BOQ Items]"] = f"{len(boq)}"
            replacements["[BOQ Total]"] = f"Rs {sum(i['amount_lac'] for i in boq):.0f} Lac"
        except Exception:
            replacements["[BOQ Items]"] = "82"
            replacements["[BOQ Total]"] = f"Rs {inv * 70:.0f} Lac"

        # Reactor dimensions
        try:
            from engines.plant_engineering import compute_all
            comp = compute_all(cfg)
            replacements["[Reactor Size]"] = f"{comp['reactor_dia_m']}m dia x {comp['reactor_ht_m']}m ht"
            replacements["[Dryer Size]"] = f"{comp['dryer_dia_m']}m dia x {comp['dryer_len_m']}m"
        except Exception:
            pass

        # Consultant
        try:
            from config import COMPANY
            replacements["[Consultant Name]"] = COMPANY.get("trade_name", "PPS Anantams")
            replacements["[Consultant Owner]"] = COMPANY.get("owner", "")
            replacements["[Consultant Phone]"] = COMPANY.get("phone", "")
            replacements["[Consultant Email]"] = COMPANY.get("email", "")
            replacements["[Consultant GST]"] = COMPANY.get("gst", "")
            replacements["[Consultant Address]"] = COMPANY.get("hq", "")
        except Exception:
            pass

    return replacements


def _replace_in_runs(paragraph, replacements):
    """Replace text in paragraph runs while preserving formatting."""
    full_text = paragraph.text
    changed = False

    for old_text, new_text in replacements.items():
        if old_text in full_text:
            full_text = full_text.replace(old_text, new_text)
            changed = True

    if changed and paragraph.runs:
        # Rebuild runs: put all text in first run, clear rest
        for i, run in enumerate(paragraph.runs):
            if i == 0:
                run.text = full_text
            else:
                run.text = ""


def customize_docx(source_path, output_path, replacements):
    """
    Open a DOCX file, replace all placeholder text, and save to output_path.
    Handles paragraphs, tables, headers, and footers.
    """
    doc = Document(source_path)

    # Replace in body paragraphs
    for paragraph in doc.paragraphs:
        _replace_in_runs(paragraph, replacements)

    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_runs(paragraph, replacements)

    # Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            if header and header.paragraphs:
                for paragraph in header.paragraphs:
                    _replace_in_runs(paragraph, replacements)
        for footer in [section.footer, section.first_page_footer]:
            if footer and footer.paragraphs:
                for paragraph in footer.paragraphs:
                    _replace_in_runs(paragraph, replacements)

    # Ensure output directory exists
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
