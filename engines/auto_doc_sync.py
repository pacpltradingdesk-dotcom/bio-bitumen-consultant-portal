"""
Auto Document Sync — When config changes, regenerate ALL documents automatically.
Fixes Gap 1 (old DOCX numbers) + Gap 2 (location not updating) + Gap 5 (govt forms).
"""
import os
import io
import json
from datetime import datetime, timezone, timedelta
from pathlib import Path

IST = timezone(timedelta(hours=5, minutes=30))
DATA_DIR = Path(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))) / "data"
SYNC_LOG = DATA_DIR / "doc_sync_log.json"


def sync_all_documents(cfg, company):
    """
    Regenerate ALL document types when config changes.
    Creates a complete package in data/auto_synced/
    """
    from engines.dynamic_doc_generator import (
        generate_dpr_docx, generate_bank_proposal_docx,
        generate_investor_pptx, generate_financial_xlsx,
    )
    from engines.report_generator_engine import generate_dpr_pdf, generate_financial_report_pdf

    output_dir = str(DATA_DIR / "auto_synced")
    os.makedirs(output_dir, exist_ok=True)

    cap = f"{cfg['capacity_tpd']:.0f}MT"
    state = cfg.get("state", "General").replace(" ", "_")
    results = {}
    errors = []

    # 1. DPR DOCX
    try:
        doc = generate_dpr_docx(cfg, company)
        path = os.path.join(output_dir, f"DPR_{cap}_{state}.docx")
        doc.save(path)
        results["DPR_DOCX"] = path
    except Exception as e:
        errors.append(f"DPR DOCX: {e}")

    # 2. Bank Proposal
    try:
        doc = generate_bank_proposal_docx(cfg, company)
        path = os.path.join(output_dir, f"Bank_Proposal_{cap}_{state}.docx")
        doc.save(path)
        results["Bank_DOCX"] = path
    except Exception as e:
        errors.append(f"Bank DOCX: {e}")

    # 3. Investor PPTX
    try:
        pptx = generate_investor_pptx(cfg, company)
        path = os.path.join(output_dir, f"Investor_Pitch_{cap}_{state}.pptx")
        pptx.save(path)
        results["Investor_PPTX"] = path
    except Exception as e:
        errors.append(f"Investor PPTX: {e}")

    # 4. Financial XLSX
    try:
        xlsx = generate_financial_xlsx(cfg, company)
        path = os.path.join(output_dir, f"Financial_Model_{cap}_{state}.xlsx")
        xlsx.save(path)
        results["Financial_XLSX"] = path
    except Exception as e:
        errors.append(f"Financial XLSX: {e}")

    # 5. DPR PDF
    try:
        path = os.path.join(output_dir, f"DPR_{cap}_{state}.pdf")
        generate_dpr_pdf(path, cfg, company)
        results["DPR_PDF"] = path
    except Exception as e:
        errors.append(f"DPR PDF: {e}")

    # 6. Financial PDF
    try:
        path = os.path.join(output_dir, f"Financial_{cap}_{state}.pdf")
        generate_financial_report_pdf(path, cfg, company)
        results["Financial_PDF"] = path
    except Exception as e:
        errors.append(f"Financial PDF: {e}")

    # 7. Government Form (auto-filled cover letter)
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("APPLICATION FOR CONSENT TO ESTABLISH", level=1)
        doc.add_paragraph(f"Date: {datetime.now(IST).strftime('%d %B %Y')}")
        doc.add_paragraph(f"To: The Member Secretary, State Pollution Control Board, {cfg.get('state', '')}")
        doc.add_paragraph(f"Subject: Application for CTE for Bio-Modified Bitumen Plant — {cfg['capacity_tpd']:.0f} MT/Day")
        doc.add_paragraph("")
        doc.add_paragraph("Respected Sir/Madam,")
        doc.add_paragraph(
            f"We, {company['trade_name']}, hereby apply for Consent to Establish (CTE) for a "
            f"Bio-Modified Bitumen manufacturing plant with a capacity of {cfg['capacity_tpd']:.0f} MT/Day "
            f"at {cfg.get('location', 'proposed location')}, {cfg.get('state', '')}."
        )
        doc.add_paragraph(f"Total Project Investment: Rs {cfg['investment_cr']:.2f} Crore")
        doc.add_paragraph(f"Employment: {cfg['staff']} persons (direct)")
        doc.add_paragraph(f"Power Requirement: {cfg['power_kw']:.0f} kW")
        doc.add_paragraph(f"Water Requirement: {max(5, int(cfg['capacity_tpd'])):,} KLD")
        doc.add_paragraph(f"Product: Bio-Modified Bitumen (VG10/VG20/VG30/VG40 grades)")
        doc.add_paragraph(f"Raw Material: Agricultural biomass (rice straw, bagasse, cotton stalk)")
        doc.add_paragraph("")
        doc.add_paragraph("Enclosed: DPR, Environmental Impact Assessment, Process Flow Diagram, Plant Layout")
        doc.add_paragraph("")
        doc.add_paragraph(f"Thanking you,")
        doc.add_paragraph(f"{company['owner']}")
        doc.add_paragraph(f"{company['trade_name']}")
        doc.add_paragraph(f"Phone: {company['phone']} | Email: {company['email']}")

        path = os.path.join(output_dir, f"Govt_CTE_Application_{cap}_{state}.docx")
        doc.save(path)
        results["Govt_CTE_DOCX"] = path
    except Exception as e:
        errors.append(f"Govt CTE: {e}")

    # 8. Factory License Application
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("APPLICATION FOR FACTORY LICENSE", level=1)
        doc.add_paragraph(f"To: The Director of Factories, {cfg.get('state', '')}")
        doc.add_paragraph(f"Date: {datetime.now(IST).strftime('%d %B %Y')}")
        doc.add_paragraph("")
        doc.add_paragraph(f"Factory Name: {company['trade_name']} — Bio-Bitumen Plant")
        doc.add_paragraph(f"Location: {cfg.get('location', '')}, {cfg.get('state', '')}")
        doc.add_paragraph(f"Nature of Manufacturing: Bio-Modified Bitumen from Agricultural Biomass Pyrolysis")
        doc.add_paragraph(f"Maximum Workers: {cfg['staff']}")
        doc.add_paragraph(f"Power (HP): {int(cfg['power_kw'] * 1.34)}")
        doc.add_paragraph(f"Investment in Plant & Machinery: Rs {cfg['investment_lac'] * 0.6:.0f} Lakhs")
        doc.add_paragraph(f"Total Project Cost: Rs {cfg['investment_cr']:.2f} Crore")
        doc.add_paragraph(f"Occupier: {company['owner']}")
        doc.add_paragraph(f"Manager: {company['owner']}")

        path = os.path.join(output_dir, f"Factory_License_App_{cap}_{state}.docx")
        doc.save(path)
        results["Factory_License_DOCX"] = path
    except Exception as e:
        errors.append(f"Factory License: {e}")

    # Log sync
    log_entry = {
        "time": datetime.now(IST).strftime("%Y-%m-%d %H:%M:%S"),
        "capacity": cfg["capacity_tpd"],
        "state": cfg.get("state", ""),
        "location": cfg.get("location", ""),
        "files_generated": len(results),
        "errors": len(errors),
        "files": list(results.keys()),
    }

    try:
        log = []
        if SYNC_LOG.exists():
            with open(str(SYNC_LOG)) as f:
                log = json.load(f)
        log.append(log_entry)
        log = log[-50:]  # Keep last 50
        with open(str(SYNC_LOG), "w") as f:
            json.dump(log, f, indent=2)
    except Exception:
        pass

    return results, errors


def get_synced_files():
    """List all auto-synced files."""
    output_dir = DATA_DIR / "auto_synced"
    if not output_dir.exists():
        return []
    files = []
    for f in sorted(output_dir.iterdir()):
        if f.is_file() and f.suffix in ('.docx', '.pdf', '.pptx', '.xlsx'):
            files.append({
                "name": f.name,
                "path": str(f),
                "size": f.stat().st_size,
                "modified": datetime.fromtimestamp(f.stat().st_mtime).strftime("%Y-%m-%d %H:%M"),
            })
    return files


def get_sync_log():
    """Get sync history."""
    try:
        if SYNC_LOG.exists():
            with open(str(SYNC_LOG)) as f:
                return json.load(f)
    except Exception:
        pass
    return []
