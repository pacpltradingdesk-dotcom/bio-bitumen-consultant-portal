"""
Bank Application Auto-Fill Engine — SBI, BOB, PNB, CGTMSE
============================================================
Generates pre-filled bank loan application forms from project data.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime


def generate_sbi_application(cfg, company):
    """Generate SBI MSME Term Loan Application (auto-filled)."""
    doc = Document()
    now = datetime.now()

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("STATE BANK OF INDIA\nMSME TERM LOAN APPLICATION")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(f"Date: {now.strftime('%d/%m/%Y')}")
    doc.add_paragraph(f"Branch: ________________________")
    doc.add_paragraph("")

    # Section 1: Applicant Details
    doc.add_heading("1. APPLICANT DETAILS", level=2)
    table = doc.add_table(rows=12, cols=2)
    table.style = 'Table Grid'
    fields = [
        ("Name of Applicant", cfg.get("client_name", "")),
        ("Constitution", cfg.get("client_company", "") + " (Pvt Ltd / LLP / Proprietorship)"),
        ("PAN Number", cfg.get("client_gst", "")[:10] if cfg.get("client_gst") else ""),
        ("GST Number", cfg.get("client_gst", "")),
        ("Address", cfg.get("site_address", "")),
        ("City / State", f"{cfg.get('location', '')} / {cfg.get('state', '')}"),
        ("Pincode", cfg.get("site_pincode", "")),
        ("Phone", cfg.get("client_phone", "")),
        ("Email", cfg.get("client_email", "")),
        ("Udyam Registration No.", "To be obtained"),
        ("Date of Incorporation", ""),
        ("Existing Banking Relationship", ""),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph("")

    # Section 2: Project Details
    doc.add_heading("2. PROJECT DETAILS", level=2)
    table2 = doc.add_table(rows=10, cols=2)
    table2.style = 'Table Grid'
    proj_fields = [
        ("Nature of Activity", "Manufacturing — Bio-Modified Bitumen"),
        ("Product", "Bio-Modified Bitumen (VG30 Grade) using Pyrolysis Technology"),
        ("Plant Capacity", f"{cfg['capacity_tpd']:.0f} MT/Day ({cfg['capacity_tpd']*300:.0f} MT/Year)"),
        ("Technology", "CSIR-CRRI Licensed Pyrolysis + VG-30 Blending"),
        ("Plant Location", f"{cfg.get('site_address', cfg.get('location', ''))} , {cfg.get('state', '')}"),
        ("Land Area", f"{cfg.get('site_area_acres', 0)} Acres"),
        ("Total Project Cost", f"Rs {cfg['investment_cr']:.2f} Crore (Rs {cfg['investment_cr']*100:.0f} Lakhs)"),
        ("Promoter's Contribution", f"Rs {cfg.get('equity_cr', 0):.2f} Crore ({cfg.get('equity_ratio', 0.4)*100:.0f}%)"),
        ("Term Loan Required", f"Rs {cfg.get('loan_cr', 0):.2f} Crore ({(1-cfg.get('equity_ratio', 0.4))*100:.0f}%)"),
        ("Working Capital Required", f"Rs {cfg.get('working_capital_lac', 0):.0f} Lakhs"),
    ]
    for i, (label, value) in enumerate(proj_fields):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = str(value)

    doc.add_paragraph("")

    # Section 3: Financial Projections
    doc.add_heading("3. FINANCIAL PROJECTIONS", level=2)
    table3 = doc.add_table(rows=8, cols=2)
    table3.style = 'Table Grid'
    fin_fields = [
        ("Annual Revenue (Year 5)", f"Rs {cfg.get('revenue_yr5_lac', 0):.0f} Lakhs"),
        ("ROI", f"{cfg.get('roi_pct', 0):.1f}%"),
        ("IRR", f"{cfg.get('irr_pct', 0):.1f}%"),
        ("DSCR (Year 3)", f"{cfg.get('dscr_yr3', 0):.2f}x"),
        ("Break-Even Period", f"{cfg.get('break_even_months', 0)} months"),
        ("Monthly EMI", f"Rs {cfg.get('emi_lac_mth', 0):.2f} Lakhs"),
        ("Interest Rate Requested", f"{cfg.get('interest_rate', 0.115)*100:.1f}% p.a."),
        ("Loan Tenure", f"{cfg.get('emi_tenure_months', 84)} months"),
    ]
    for i, (label, value) in enumerate(fin_fields):
        table3.rows[i].cells[0].text = label
        table3.rows[i].cells[1].text = str(value)

    doc.add_paragraph("")

    # Section 4: Security Offered
    doc.add_heading("4. SECURITY OFFERED", level=2)
    doc.add_paragraph("Primary Security: Hypothecation of Plant & Machinery, Equipment")
    doc.add_paragraph("Collateral Security: ________________________ (Land / Property / FD)")
    doc.add_paragraph(f"CGTMSE: {'Eligible (Loan < Rs 5 Crore)' if cfg.get('loan_cr', 0) <= 5 else 'Not Eligible (Loan > Rs 5 Crore) — Collateral Required'}")

    doc.add_paragraph("")

    # Section 5: Documents Enclosed
    doc.add_heading("5. DOCUMENTS ENCLOSED", level=2)
    checklist = [
        "Detailed Project Report (DPR)",
        "CMA Data (7-Year Projections)",
        "PAN Card / Aadhaar of Promoter",
        "GST Registration Certificate",
        "Udyam (MSME) Registration",
        "Property documents for collateral",
        "Last 3 years ITR (if existing business)",
        "Bank statements (6 months)",
        "Quotations from machinery suppliers",
        "Plant layout and drawings",
    ]
    for item in checklist:
        doc.add_paragraph(f"☐ {item}")

    doc.add_paragraph("")
    doc.add_paragraph("")

    # Signature
    doc.add_paragraph("_________________________")
    doc.add_paragraph(f"Signature of Applicant")
    doc.add_paragraph(f"{cfg.get('client_name', '')}")
    doc.add_paragraph(f"Date: {now.strftime('%d/%m/%Y')}")

    doc.add_paragraph("")
    doc.add_paragraph(f"Prepared with assistance from: {company.get('trade_name', 'PPS Anantams')}")
    doc.add_paragraph(f"Consultant: {company.get('owner', '')} | {company.get('phone', '')}")

    return doc


def generate_cgtmse_application(cfg, company):
    """Generate CGTMSE Guarantee Application Form."""
    doc = Document()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CREDIT GUARANTEE FUND TRUST FOR MICRO AND SMALL ENTERPRISES (CGTMSE)\nAPPLICATION FOR GUARANTEE COVER")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(f"Date: {datetime.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph("")

    doc.add_heading("BORROWER DETAILS", level=2)
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    fields = [
        ("Name of Borrower", cfg.get("client_name", "")),
        ("Enterprise Name", cfg.get("client_company", "")),
        ("Constitution", "Private Limited / LLP"),
        ("Udyam Registration", "Applied / Obtained"),
        ("Category", "Micro / Small Enterprise"),
        ("Activity", "Manufacturing — Bio-Modified Bitumen"),
        ("Address", f"{cfg.get('site_address', '')} , {cfg.get('location', '')} , {cfg.get('state', '')}"),
        ("Bank Branch", "________________________"),
    ]
    for i, (label, value) in enumerate(fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = str(value)

    doc.add_paragraph("")
    doc.add_heading("LOAN DETAILS", level=2)

    loan_cr = cfg.get("loan_cr", 0)
    eligible = loan_cr <= 5

    table2 = doc.add_table(rows=6, cols=2)
    table2.style = 'Table Grid'
    loan_fields = [
        ("Total Loan Amount", f"Rs {loan_cr:.2f} Crore"),
        ("CGTMSE Eligible Amount", f"Rs {min(loan_cr, 5):.2f} Crore" if eligible else "Rs 5.00 Crore (Max limit)"),
        ("Guarantee Fee (Annual)", f"1.5% = Rs {min(loan_cr, 5)*1.5:.2f} Lakhs/year"),
        ("Collateral Required", "NIL (CGTMSE Guarantee)" if eligible else f"Rs {loan_cr-5:.2f} Crore portion requires collateral"),
        ("Loan Tenure", f"{cfg.get('emi_tenure_months', 84)} months"),
        ("Purpose", "Term Loan for Bio-Modified Bitumen Plant Setup"),
    ]
    for i, (label, value) in enumerate(loan_fields):
        table2.rows[i].cells[0].text = label
        table2.rows[i].cells[1].text = str(value)

    doc.add_paragraph("")
    doc.add_paragraph("_________________________")
    doc.add_paragraph("Authorized Signatory of Lending Bank")

    return doc


BANK_FORMS = {
    "SBI MSME Term Loan": generate_sbi_application,
    "CGTMSE Guarantee": generate_cgtmse_application,
}
